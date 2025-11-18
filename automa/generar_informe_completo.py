import os
import time
import sys
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx2pdf import convert
from PyPDF2 import PdfMerger

sys.stdout.reconfigure(encoding='utf-8')

# =======================
# CONFIGURACIÓN GENERAL
# =======================
carpeta_capturas = r"C:\Users\Edinson\Desktop\TODO\CODIGOS\14_CONSULTOR\V7\seleniumweb\media\descargas"

nombre_base = "INFORME DE CONSULTAS"

salida_word = os.path.join(carpeta_capturas, f"{nombre_base}.docx")
salida_pdf = os.path.join(carpeta_capturas, f"{nombre_base}.pdf")


# =======================
# FUNCIÓN PARA GENERAR WORD Y PDF
# =======================
def generar_word_y_pdf(carpeta):
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    doc.add_heading("INFORME DE CONSULTAS", level=1)

    # Orden especificado por el usuario
    orden_capturas = [
        "ofac_final.png",
        "contaduria_final.png",
        "rues_evento_2.png",
        "rues_evento_5.png",
        "rues_evento_7.png",
        "rues_evento_9.png",
        "rues_evento_10.png",
        "rues_final.png"
    ]

    for i, archivo in enumerate(orden_capturas):
        img_path = os.path.join(carpeta, archivo)

        if not os.path.exists(img_path):
            continue

        nombre_titulo = (
            os.path.splitext(archivo)[0]
            .replace("_final", "")
            .replace("_", " ")
            .upper()
        )

        doc.add_heading(nombre_titulo, level=2)

        try:
            doc.add_picture(img_path, width=Inches(10))
        except Exception as e:
            print(f"[ERROR] No se pudo insertar {archivo}: {e}")

        if i < len(orden_capturas) - 1:
            doc.add_page_break()

    doc.save(salida_word)
    print(f"[OK] Word generado: {salida_word}")

    try:
        convert(salida_word, salida_pdf)
        print(f"[OK] PDF generado: {salida_pdf}")
    except Exception as e:
        print(f"[ERROR] Conversión a PDF falló: {e}")
        return None

    # Esperar creación real del PDF
    for _ in range(10):
        if os.path.exists(salida_pdf) and os.path.getsize(salida_pdf) > 1000:
            break
        time.sleep(1)

    return salida_pdf


# =======================
# FUNCIÓN PARA COMBINAR PDFs
# =======================
def combinar_pdfs(carpeta, pdf_base):

    pdfs_numericos = []
    numeros_detectados = []

    for archivo in os.listdir(carpeta):

        # Solo considerar archivos PDF
        if not archivo.lower().endswith(".pdf"):
            continue

        # Ignorar el PDF temporal base
        if archivo == f"{nombre_base}.pdf":
            continue

        # Ignorar informes anteriores
        if archivo.startswith(f"{nombre_base}_"):
            continue

        # Extraer dígitos iniciales
        nombre = os.path.splitext(archivo)[0]
        num = ""

        for c in nombre:
            if c.isdigit():
                num += c
            else:
                break

        # Guardar PDFs numéricos válidos
        if num.isdigit():
            pdfs_numericos.append(os.path.join(carpeta, archivo))
            numeros_detectados.append(int(num))

    # Si no hay PDFs numéricos → salir
    if not pdfs_numericos:
        print("No se encontraron PDFs numéricos.")
        return None, None

    # Ordenar los PDFs correctamente
    def extraer_numero(path):
        nombre = os.path.splitext(os.path.basename(path))[0]
        num = ""
        for c in nombre:
            if c.isdigit():
                num += c
            else:
                break
        return int(num)

    pdfs_numericos.sort(key=extraer_numero)

    # Primer número detectado
    numero_pdf = numeros_detectados[0]

    # Nombre del informe final
    salida_final = os.path.join(carpeta, f"{nombre_base}_{numero_pdf}.pdf")

    # Unir PDFs
    merger = PdfMerger()
    merger.append(pdf_base)

    for pdf in pdfs_numericos:
        merger.append(pdf)

    merger.write(salida_final)
    merger.close()

    print(f"[OK] PDF final generado: {salida_final}")

    return salida_final, pdfs_numericos


# =======================
# FUNCIÓN PARA LIMPIAR ARCHIVOS SOBRANTES
# =======================
def limpiar_archivos(carpeta, archivo_final, pdfs_numericos):
    for archivo in os.listdir(carpeta):
        ruta = os.path.join(carpeta, archivo)

        # Mantener SIEMPRE el PDF final generado
        if ruta == archivo_final:
            continue

        # ============================
        # PROTECCIÓN: NO BORRAR INFORMES ANTERIORES
        # ============================
        # Mantener archivos tipo: INFORME DE CONSULTAS_12345.pdf
        if archivo.startswith(f"{nombre_base}_") and archivo.endswith(".pdf"):

            parte = archivo.replace(f"{nombre_base}_", "").replace(".pdf", "")

            if parte.isdigit():
                continue  # → Es un informe previo, NO BORRAR

        # ============================
        # ARCHIVOS QUE SÍ SE DEBEN BORRAR
        # ============================

        # Archivo Word temporal
        if archivo.endswith(".docx"):
            os.remove(ruta)
            continue

        # Imagen temporal
        if archivo.lower().endswith((".png", ".jpg", ".jpeg")):
            os.remove(ruta)
            continue

        # PDF base temporal
        if archivo == f"{nombre_base}.pdf":
            os.remove(ruta)
            continue

        # PDFs numéricos usados para combinar
        if ruta in pdfs_numericos:
            os.remove(ruta)
            continue

        # Cualquier otro PDF que no sea final ni informe antiguo
        if archivo.endswith(".pdf") and ruta != archivo_final:
            os.remove(ruta)

    print("[OK] Archivos sobrantes eliminados.")


# =======================
# MAIN
# =======================
if __name__ == "__main__":
    if not os.path.exists(carpeta_capturas):
        print("La carpeta no existe.")
    else:
        pdf_base = generar_word_y_pdf(carpeta_capturas)

        if pdf_base:
            archivo_final, pdfs_num = combinar_pdfs(carpeta_capturas, pdf_base)

            if archivo_final:
                limpiar_archivos(carpeta_capturas, archivo_final, pdfs_num)
