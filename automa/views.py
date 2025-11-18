from django.http import JsonResponse, FileResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from django.shortcuts import render
from django.conf import settings
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
import os
import time
import tempfile
import pypandoc
import traceback
import subprocess
import pathlib
import shutil

from .selenium_script import ejecutar_consulta


# ==========================================================
# 🖼️ Capturas específicas que se incluirán en el informe
# ==========================================================
CAPTURAS_INFORME = [
    "ofac_final.png",
    "contaduria_final.png",
    "ofac_final.png"
]


# ==========================================================
# 🔹 API: Ejecutar Selenium y generar PDF automáticamente
# ==========================================================
@csrf_exempt
def run_consulta(request):
    """
    Ejecuta el script Selenium, selecciona capturas específicas,
    genera un informe PDF y devuelve su URL pública.
    """
    if request.method != "POST":
        return JsonResponse({
            "status": "error",
            "msg": "Método no permitido (usa POST)"
        }, status=405)

    numero_doc = request.POST.get("numero")
    if not numero_doc:
        return JsonResponse({
            "status": "error",
            "msg": "⚠️ Falta el número de documento"
        }, status=400)

    try:
        inicio = time.time()
        resultado = ejecutar_consulta(numero_doc)
        duracion = time.time() - inicio

        # 🔁 Filtrar solo las capturas deseadas que existan
        capturas_seleccionadas = [
            c for c in resultado.get("capturas", [])
            if os.path.basename(c) in CAPTURAS_INFORME and os.path.exists(c)
        ]

        # 🧾 Crear documento Word temporal
        doc = Document()

        # Cambiar orientación a horizontal (landscape)
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        # Título principal
        doc.add_heading(f"Informe de Consulta - {numero_doc}", level=1)
        doc.add_paragraph(f"Duración: {duracion:.2f} segundos")

        # Solo agregar título de capturas si existen
        if capturas_seleccionadas:
            doc.add_paragraph("Capturas incluidas en este informe:\n")

            for ruta in capturas_seleccionadas:
                nombre = os.path.basename(ruta)
                doc.add_paragraph(nombre)
                doc.add_picture(ruta, width=Inches(6.5))

        # 📄 Guardar DOCX temporalmente
        tmp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(tmp_dir, "informe.docx")
        pdf_path = os.path.join(tmp_dir, "informe.pdf")
        doc.save(docx_path)

        # 🔄 Convertir DOCX → PDF con pypandoc y wkhtmltopdf
        try:
            wkhtmltopdf_path = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"

            # 1️⃣ DOCX → HTML
            html_path = os.path.join(tmp_dir, "informe.html")
            pypandoc.convert_file(docx_path, "html", outputfile=html_path, extra_args=["--standalone"])

            # 2️⃣ HTML → PDF con wkhtmltopdf
            result = subprocess.run(
                [wkhtmltopdf_path, html_path, pdf_path],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True
                )
            if result.returncode != 0:
                raise Exception(f"wkhtmltopdf falló: {result.stderr}")

        except Exception as e:
            print("🧨 ERROR COMPLETO EN run_consulta():")
            print(traceback.format_exc())
            return JsonResponse({
                "status": "error",
                "msg": f"❌ Error generando PDF: {str(e)}"
            }, status=500)

        # 📂 Guardar el PDF final en /media/descargas/
        carpeta_salida = os.path.join(settings.MEDIA_ROOT, "descargas")
        os.makedirs(carpeta_salida, exist_ok=True)
        destino = os.path.join(carpeta_salida, f"informe_{numero_doc}.pdf")
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"No se generó el PDF en {pdf_path}")
        shutil.move(pdf_path, destino)

        # 🌐 URL pública del PDF generado
        url_pdf = settings.MEDIA_URL + f"descargas/informe_{numero_doc}.pdf"

        return JsonResponse({
            "status": "ok",
            "msg": f"✅ Consulta completada para {numero_doc}",
            "tiempo": f"{duracion:.2f} segundos",
            "informe_pdf": url_pdf
        })

    except Exception as e:
        print("🧨 ERROR EN run_consulta():", traceback.format_exc())
        return JsonResponse({
            "status": "error",
            "msg": f"❌ Error en la consulta: {str(e)}"
        }, status=500)


# ==========================================================
# 🔹 Página principal
# ==========================================================
def index(request):
    """Renderiza la página principal con el formulario."""
    return render(request, "automa/index.html")


# ==========================================================
# 🔹 Listar archivos en /media/descargas/
# ==========================================================
def listar_archivos(request):
    carpeta_descargas = os.path.join(settings.MEDIA_ROOT, "descargas")

    if not os.path.exists(carpeta_descargas):
        return render(request, "automa/listar.html", {
            "archivos": [],
            "error": "⚠️ La carpeta 'descargas' no existe."
        })

    archivos = []
    for nombre in os.listdir(carpeta_descargas):
        ruta_completa = os.path.join(carpeta_descargas, nombre)
        if os.path.isfile(ruta_completa):
            archivos.append({
                "nombre": nombre,
                "url": f"{settings.MEDIA_URL}descargas/{nombre}"
            })

    filtro = request.GET.get("filtro")
    if filtro:
        archivos = [a for a in archivos if filtro.lower() in a["nombre"].lower()]

    return render(request, "automa/listar.html", {"archivos": archivos})


# ==========================================================
# 🔹 Eliminar archivos seleccionados
# ==========================================================
@require_POST
def eliminar_archivos(request):
    seleccionados = request.POST.getlist("archivos")
    carpeta_descargas = os.path.join(settings.MEDIA_ROOT, "descargas")

    eliminados, errores = [], []

    for nombre in seleccionados:
        ruta = os.path.join(carpeta_descargas, nombre)
        try:
            os.remove(ruta)
            eliminados.append(nombre)
        except Exception as e:
            errores.append(f"{nombre}: {str(e)}")

    mensaje = f"✅ {len(eliminados)} archivo(s) eliminado(s)."
    if errores:
        mensaje += f" ⚠️ {len(errores)} no se pudieron eliminar."

    archivos = []
    for nombre in os.listdir(carpeta_descargas):
        ruta_completa = os.path.join(carpeta_descargas, nombre)
        if os.path.isfile(ruta_completa):
            archivos.append({
                "nombre": nombre,
                "url": f"{settings.MEDIA_URL}descargas/{nombre}"
            })

    return render(request, "automa/listar.html", {
        "archivos": archivos,
        "mensaje": mensaje
    })


# ==========================================================
# 🔹 Generar y descargar PDF combinado (manual)
# ==========================================================
def generar_y_descargar_pdf(request):
    """
    Ejecuta el script Python que genera el informe completo y
    devuelve el archivo PDF final al navegador.
    """
    carpeta_capturas = os.path.join(settings.BASE_DIR, "media", "descargas")
    script_path = os.path.join(settings.BASE_DIR, "automa", "generar_informe_completo.py")

    try:
        # Ejecutar script externo
        proceso = subprocess.run(
            ["python", str(script_path)],
            capture_output=True,
            text=True
        )

        print("=== SALIDA SCRIPT ===")
        print(proceso.stdout)
        print(proceso.stderr)

        if proceso.returncode != 0:
            return HttpResponse(
                f"⚠️ Error ejecutando el script:<br><pre>{proceso.stderr}</pre>",
                content_type="text/html",
                status=500
            )

        # Buscar PDF más reciente
        pdfs = [f for f in os.listdir(carpeta_capturas) if f.lower().endswith(".pdf")]
        if not pdfs:
            return HttpResponse("⚠️ No se encontró ningún PDF generado.", status=404)

        pdfs.sort(key=lambda f: os.path.getmtime(os.path.join(carpeta_capturas, f)), reverse=True)
        pdf_final = os.path.join(carpeta_capturas, pdfs[0])

        return FileResponse(open(pdf_final, "rb"), as_attachment=True, filename=os.path.basename(pdf_final))

    except Exception as e:
        error_trace = traceback.format_exc()
        print("🧨 ERROR en generar_y_descargar_pdf:\n", error_trace)
        return HttpResponse(
            f"❌ Error general al generar o descargar el informe:<br><pre>{error_trace}</pre>",
            content_type="text/html",
            status=500
        )
from django.http import JsonResponse

def listar_archivos_json(request):
    carpeta = os.path.join(settings.MEDIA_ROOT, "descargas")

    archivos = []

    if os.path.exists(carpeta):
        for nombre in os.listdir(carpeta):
            if nombre.lower().endswith(".pdf"):
                archivos.append({
                    "nombre": nombre,
                    "url": settings.MEDIA_URL + "descargas/" + nombre
                })

    return JsonResponse({"archivos": archivos})
